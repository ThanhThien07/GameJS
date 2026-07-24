import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def build_perfect_game_report():
    sys.stdout.reconfigure(encoding='utf-8')

    templates = [f for f in os.listdir('.') if f.endswith('.docx') and 'BaoCao_GameClicker' not in f]
    if templates:
        template_file = templates[0]
    else:
        template_file = 'BaoCao_GameClicker.docx'

    print(f"Loading cover template from: {template_file}")
    doc = docx.Document(template_file)

    # 1. SET PRINT-READY MARGINS (Left: 3.0 cm, Right: 2.0 cm, Top: 2.0 cm, Bottom: 2.0 cm)
    for section in doc.sections:
        section.top_margin = Inches(0.79)    # ~2.0 cm
        section.bottom_margin = Inches(0.79) # ~2.0 cm
        section.left_margin = Inches(1.18)   # ~3.0 cm for binding
        section.right_margin = Inches(0.79)  # ~2.0 cm

    # 2. UPDATE COVER PAGE (P0 to P23) EXACTLY AS TEMPLATE FORMAT
    def set_p_cover(p, text, bold=True, size_pt=14):
        align = p.alignment
        p.text = text
        p.alignment = align
        if len(p.runs) > 0:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.bold = bold
                r.font.color.rgb = RGBColor(0, 0, 0)
                if size_pt:
                    r.font.size = Pt(size_pt)

    # P7: Title
    set_p_cover(doc.paragraphs[7], "PHÂN TÍCH THIẾT KẾ VÀ XÂY DỰNG GAME SIÊU CLICKER TAM HỢP\n(TAP TAP CLICKER MULTIPLAYER)", bold=True, size_pt=18)
    
    # P8: Semester
    set_p_cover(doc.paragraphs[8], "HỌC KỲ 3, NĂM HỌC 2025 - 2026", bold=True, size_pt=14)
    
    # P10: Group Name
    set_p_cover(doc.paragraphs[10], "NHÓM DECIBEL", bold=True, size_pt=16)

    # P15: NSVTH 1
    set_p_cover(doc.paragraphs[15], "NSVTH: NGUYỄN HOÀNG HƯNG  (MSSV: 501250384)", bold=True, size_pt=13)
    
    # P16: NSVTH 2
    set_p_cover(doc.paragraphs[16], "\t\t   HOÀNG THỊ DIỄM LY", bold=True, size_pt=13)
    
    # P17: Clear 3rd author
    set_p_cover(doc.paragraphs[17], "", bold=False, size_pt=13)

    # P18: Subject
    set_p_cover(doc.paragraphs[18], "MÔN : JAVASCRIPTS", bold=True, size_pt=13)
    
    # P19: Class
    set_p_cover(doc.paragraphs[19], "LỚP: CD25CT6", bold=True, size_pt=13)

    # P20: Lecturer
    set_p_cover(doc.paragraphs[20], "GVHD: ThS. Trương Châu Long", bold=True, size_pt=13)

    # P23: Date
    set_p_cover(doc.paragraphs[23], "Tp. Hồ Chí Minh, tháng 07 năm 2026", bold=True, size_pt=12)

    # Remove ALL paragraphs after index 23 (removes duplicate title block, old food content, etc.)
    p_elements = [p._element for p in doc.paragraphs[24:]]
    for pe in p_elements:
        pe.getparent().remove(pe)

    # Remove old tables
    t_elements = [t._element for t in doc.tables]
    for te in t_elements:
        te.getparent().remove(te)

    # Helper function for setting cell background
    def set_cell_bg(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Print-ready paragraph helper (Times New Roman, Justified, Black, 1.2 Line Spacing, First Line Indent 0.5 in)
    def add_body_p(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, size=13, color=(0, 0, 0), space_before=2, space_after=4, line_spacing=1.2):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.first_line_indent = Inches(0.5) # Standard academic indent
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_bullet(text, bold_prefix="", level=0):
        p = doc.add_paragraph(style='List Paragraph')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
        
        run_bullet = p.add_run("•  ")
        run_bullet.font.name = 'Times New Roman'
        run_bullet.font.size = Pt(13)
        run_bullet.font.bold = True
        run_bullet.font.color.rgb = RGBColor(0, 0, 0)

        if bold_prefix:
            run_b = p.add_run(bold_prefix + " ")
            run_b.font.name = 'Times New Roman'
            run_b.font.size = Pt(13)
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(0, 0, 0)

        run_t = p.add_run(text)
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(13)
        run_t.font.color.rgb = RGBColor(0, 0, 0)
        return p

    # PAGE BREAK -> LỜI CẢM ƠN (Page 2)
    doc.add_page_break()

    # ==========================================
    # LỜI CẢM ƠN
    # ==========================================
    add_h1("LỜI CẢM ƠN")
    add_body_p("Lời đầu tiên, nhóm chúng em xin gửi lời cảm ơn chân thành và sâu sắc nhất đến Ban Giám hiệu Trường Cao đẳng Công nghệ Thông tin TP. Hồ Chí Minh cùng toàn thể quý Thầy Cô Khoa Công nghệ Thông tin – Điện tử đã tạo điều kiện học tập tốt nhất, môi trường rèn luyện hiện đại và truyền đạt những nền tảng kiến thức chuyên môn quý báu cho chúng em trong suốt thời gian qua.")
    add_body_p("Đặc biệt, nhóm chúng em xin bày tỏ lòng biết ơn sâu sắc nhất tới thầy ThS. Trương Châu Long — giảng viên phụ trách môn JAVASCRIPTS. Thầy đã tận tình hướng dẫn, trang bị cho chúng em những tư duy lập trình hiện đại, phương pháp xây dựng ứng dụng Web thời gian thực và luôn đưa ra những lời khuyên, đóng góp định hướng kịp thời để nhóm hoàn thiện đồ án này một cách chỉn chu và bài bản nhất.")
    add_body_p("Mặc dù nhóm đã dành nhiều tâm huyết và nỗ lực để hoàn thành đồ án, song do kiến thức và kinh nghiệm thực tế của chúng em còn hạn chế nên sản phẩm khó tránh khỏi những thiếu sót. Nhóm rất mong nhận được sự chỉ bảo, nhận xét và đánh giá quý báu từ thầy ThS. Trương Châu Long để sản phẩm của chúng em ngày càng hoàn thiện hơn.")
    add_body_p("Nhóm xin trân trọng và chân thành cảm ơn!")

    # PAGE BREAK -> MỤC LỤC (Page 3)
    doc.add_page_break()

    # ==========================================
    # MỤC LỤC (Formatted with Word Native Dot Leaders)
    # ==========================================
    add_h1("MỤC LỤC")
    
    toc_items = [
        ("LỜI CẢM ƠN", "2"),
        ("MỤC LỤC", "3"),
        ("MỞ ĐẦU", "4"),
        ("CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI GAME CLICKER", "5"),
        ("    1.1. Bối cảnh và Tính cấp thiết của Đề tài", "5"),
        ("    1.2. Mục tiêu nghiên cứu và Xây dựng Game", "5"),
        ("    1.3. Lợi ích và Phạm vi ứng dụng của Sản phẩm", "6"),
        ("CHƯƠNG 2: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG GAME", "7"),
        ("    2.1. Phân hệ Chế độ Chơi Ngoại tuyến (Offline Mode)", "7"),
        ("    2.2. Phân hệ Chế độ Chơi Đấu trường Trực tuyến (Competitive Real-time 1v1v1)", "8"),
        ("    2.3. Phân hệ Chế độ Chơi Hợp tác Trực tuyến (Co-op 3 Players)", "9"),
        ("    2.4. Phân hệ Đồ họa 3D Cartoon Theme & Web Audio Synthesizer", "10"),
        ("    2.5. Phân hệ Quản lý Cơ sở Dữ liệu MySQL & Server Cloud", "10"),
        ("CHƯƠNG 3: ĐẶC TẢ SƠ ĐỒ LUỒNG XỬ LÝ (WORKFLOWS)", "11"),
        ("    3.1. Luồng xử lý của Người chơi Offline", "11"),
        ("    3.2. Luồng xử lý của Người chơi Online Đấu trường", "11"),
        ("    3.3. Luồng xử lý của Người chơi Online Hợp tác", "12"),
        ("CHƯƠNG 4: THIẾT KẾ CƠ SỞ DỮ LIỆU VÀ MÁY CHỦ MULTIPLAYER", "13"),
        ("    4.1. Cấu trúc Thực thể Cơ sở Dữ liệu (Database Schema)", "13"),
        ("    4.2. Đặc tả Chi tiết các Bảng CSDL Cốt lõi", "13"),
        ("    4.3. Xử lý các Tình huống Vận hành Thực tế", "14"),
        ("CHƯƠNG 5: THIẾT LẬP CÔNG NGHỆ VÀ TRIỂN KHAI CLOUD", "15"),
        ("    5.1. Công nghệ Frontend UI Engine", "15"),
        ("    5.2. Công nghệ Backend Engine & WebSockets", "15"),
        ("    5.3. Môi trường Triển khai Production Cloud", "16"),
        ("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "17")
    ]

    for title, pg in toc_items:
        p_toc = doc.add_paragraph()
        p_toc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_toc.paragraph_format.space_before = Pt(2)
        p_toc.paragraph_format.space_after = Pt(3)
        p_toc.paragraph_format.line_spacing = 1.15
        
        # Native Word Tab with Dot Leader at 6.3 in (right margin)
        pPr = p_toc._p.get_or_add_pPr()
        tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:leader="dot" w:pos="9072"/></w:tabs>')
        pPr.append(tabs)

        is_main = title.startswith("CHƯƠNG") or title.strip() in ["LỜI CẢM ƠN", "MỤC LỤC", "MỞ ĐẦU", "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN"]
        
        run_t = p_toc.add_run(title)
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(12)
        run_t.font.bold = is_main
        run_t.font.color.rgb = RGBColor(0, 0, 0)

        run_tab = p_toc.add_run("\t")
        run_tab.font.name = 'Times New Roman'

        run_p = p_toc.add_run(pg)
        run_p.font.name = 'Times New Roman'
        run_p.font.size = Pt(12)
        run_p.font.bold = is_main
        run_p.font.color.rgb = RGBColor(0, 0, 0)

    # PAGE BREAK -> MỞ ĐẦU (Page 4)
    doc.add_page_break()

    # ==========================================
    # MỞ ĐẦU
    # ==========================================
    add_h1("MỞ ĐẦU")
    add_body_p("Trong kỷ nguyên số hóa hiện nay, ngành công nghiệp giải trí kỹ thuật số và phát triển Game Web trên nền tảng JavaScript đang bùng nổ mạnh mẽ nhờ tính linh hoạt, khả năng tương thích cao trên mọi thiết bị di động cũng như máy tính mà không cần cài đặt phức tạp. Dòng game Clicker (Idle Clicker Game) từ lâu đã trở thành một hiện tượng thu hút hàng triệu người chơi nhờ lối chơi đơn giản nhưng cực kỳ cuốn hút, giúp giải tỏa căng thẳng nhanh chóng.")
    add_body_p("Tuy nhiên, phần lớn các tựa game Clicker truyền thống hiện nay chỉ dừng lại ở chế độ chơi đơn (Offline) thụ động, thiếu tính kết nối cộng đồng và dễ gây cảm giác nhàm chán sau một thời gian ngắn. Xuất phát từ nhu cầu thực tiễn đó, nhóm em đã lựa chọn đề tài \"Phân tích thiết kế và xây dựng Game Siêu Clicker Tam Hợp (Tap Tap Clicker Multiplayer)\" để tạo ra một sản phẩm game Web hiện đại, kết hợp hài hòa giữa tính năng chơi đơn tích lũy và chế độ thi đấu/hợp tác trực tuyến đa người chơi (Real-time Multiplayer) thông qua công nghệ WebSocket tiên tiến.")

    # ==========================================
    # CHƯƠNG 1
    # ==========================================
    add_h1("CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI GAME CLICKER")
    
    add_h2("1.1. Bối cảnh và Tính cấp thiết của Đề tài")
    add_body_p("Ngành game Web hiện đại đòi hỏi sự đổi mới không ngừng về trải nghiệm người dùng. Việc kết hợp cơ chế Click tích lũy điểm số truyền thống với các công nghệ giao tiếp thời gian thực đem lại sự đột phá lớn trong tương tác nhóm.")
    add_bullet("Các game Web Clicker thông thường chỉ cho phép người chơi tương tác một mình, thiếu tính cạnh tranh và giao lưu.", "Hạn chế của game truyền thống:")
    add_bullet("Ứng dụng bộ công cụ tiên tiến React 19, Vite, Node.js, Socket.io và cơ sở dữ liệu MySQL để xây dựng hạ tầng game đa người chơi hoàn chỉnh.", "Ứng dụng công nghệ JavaScript mới:")

    add_h2("1.2. Mục tiêu nghiên cứu và Xây dựng Game")
    add_bullet("Thiết kế giao diện người dùng đẹp mắt, hiện đại theo phong cách Glassmorphism với hiệu ứng vầng ánh sáng mặt trời xoay (Sunburst Rays).", "Xây dựng giao diện đồ họa sống động:")
    add_bullet("Phát triển 3 chế độ chơi gồm: Ngoại tuyến (Offline), Trực tuyến Đấu trường (Competitive 1v1v1) và Trực tuyến Hợp tác (Co-op 3 người).", "Đa dạng chế độ chơi:")
    add_bullet("Xây dựng bộ tổng hợp âm thanh Web Audio API thuần bằng code mà không cần tải các file âm thanh bên ngoài.", "Tối ưu hóa âm thanh Web Audio API:")

    add_h2("1.3. Lợi ích và Phạm vi ứng dụng của Sản phẩm")
    add_body_p("Sản phẩm đáp ứng nhu cầu giải trí nhanh trên mọi nền tảng trình duyệt Web. Người chơi có thể lưu trữ tiến trình chơi cá nhân mượt mà hoặc kết nối thi đấu trực tiếp với bạn bè qua mạng Internet mà không gặp rào cản cài đặt.")

    # ==========================================
    # CHƯƠNG 2
    # ==========================================
    add_h1("CHƯƠNG 2: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG GAME")
    add_body_p("Hệ thống Siêu Clicker Tam Hợp được phân chia thành các phân hệ chức năng độc lập, hỗ trợ linh hoạt cả hai chế độ chơi Ngoại tuyến và Trực tuyến:")

    add_h2("2.1. Phân hệ Chế độ Chơi Ngoại tuyến (Offline Mode)")
    add_bullet("Người chơi nhấp chuột vào linh vật 3D để tạo điểm và tăng chỉ số DPC (Damage Per Click).", "Cơ chế Click tay (Manual Tapping):")
    add_bullet("Thuê các công cụ hỗ trợ tự động (Găng tay sắt, Rìu chặt củi, Kiếm kim cương, Xe goòng mỏ, Máy khoan laze, Giàn khoan) để tăng nguồn thu nhập DPS (Damage Per Second).", "Hệ thống Nâng cấp tự động (Auto Helpers):")
    add_bullet("Khi đạt 50,000 vàng, người chơi thực hiện Trùng Sinh để đổi lấy Tinh Thể Linh Hồn (Soul Crystals), gia tăng vĩnh viễn +15% chỉ số.", "Hệ thống Trùng Sinh (Prestige/Rebirth):")
    add_bullet("Kích hoạt Cơn Cuồng Phong (x2 DPC trong 10s) và Bão Vàng (Thưởng vàng lập tức).", "Bộ Kỹ năng Nguồn lực (Active Skills):")

    add_h2("2.2. Phân hệ Chế độ Chơi Đấu trường Trực tuyến (Competitive Real-time 1v1v1)")
    add_bullet("Hệ thống phát sinh mã phòng ngẫu nhiên 6 ký tự để người chơi tạo phòng hoặc tham gia phòng đấu.", "Tạo và Gia nhập Phòng (Room Management):")
    add_bullet("Cho phép chủ phòng thêm các đối thủ máy (AI Bots) với tần số nhấp chuột tự động ngẫu nhiên khi thiếu người.", "Tích hợp AI Bot tự động:")
    add_bullet("Đếm ngược 30 giây thi đấu nghẹt thở. Điểm số và nhịp click được đồng bộ thời gian thực qua WebSocket.", "Đấu trường 30 giây Real-time:")
    add_bullet("Kết xuất bảng vàng vinh danh thứ hạng #1, #2, #3 ngay sau khi kết thúc trận đấu.", "Bảng xếp hạng Trận đấu:")

    add_h2("2.3. Phân hệ Chế độ Chơi Hợp tác Trực tuyến (Co-op 3 Players)")
    add_bullet("3 người chơi chung tay nhấp chuột hạ gục Mục tiêu Boss/Cổ thụ/Quặng đá.", "Đồng lòng hạ gục Mục tiêu:")
    add_bullet("Các vật phẩm Thịt 🥩, Gỗ 🪵, Đá 🪨 rơi ra ngẫu nhiên khi nhấp chuột được đồng bộ lập tức cho mọi người trong phòng.", "Rơi tài nguyên ngẫu nhiên (Resource Drops):")
    add_bullet("Người chơi đóng góp tài nguyên thu gom được để mua Nâng cấp Sát thương chung, Hệ số nhân và Robot Auto-Click cho toàn phòng.", "Nâng cấp dùng chung (Shared Upgrades):")

    add_h2("2.4. Phân hệ Đồ họa 3D Cartoon Theme & Web Audio Synthesizer")
    add_bullet("Đánh Quái Vật (Monster Fight), Tiều Phu Chặt Gỗ (Woodcutter), Thợ Mỏ Đào Đá (Stone Mining).", "3 Đồ họa Theme 3D Hoạt hình:")
    add_bullet("Tự động tạo các dải sóng âm thanh (Click, Buy, Skill, Rebirth) trực tiếp qua Web Audio API.", "Bộ âm thanh Web Audio API:")

    add_h2("2.5. Phân hệ Quản lý Cơ sở Dữ liệu MySQL & Server Cloud")
    add_bullet("Lưu giữ kết quả các phòng đấu, điểm số người chơi và lịch sử trận đấu.", "Cơ sở dữ liệu MySQL:")
    add_bullet("Triển khai trên máy chủ Railway Cloud với listener binding `0.0.0.0:${PORT}` bảo đảm kết nối mượt mà 24/7.", "Triển khai Cloud Server Railway:")

    # ==========================================
    # CHƯƠNG 3
    # ==========================================
    add_h1("CHƯƠNG 3: ĐẶC TẢ SƠ ĐỒ LUỒNG XỬ LÝ (WORKFLOWS)")

    add_h2("3.1. Sơ đồ Luồng xử lý Chơi đơn Offline")
    add_body_p("Người chơi mở Game → Chọn Chế độ Ngoại tuyến → Chọn Mô hình Đồ họa Theme → Nhấp chuột tích điểm → Thuê công cụ nâng cấp DPC/DPS → Tích lũy sạc Nộ x2 → Kích hoạt Kỹ năng / Trùng sinh → Dữ liệu tự động lưu vào LocalStorage.")

    add_h2("3.2. Sơ đồ Luồng xử lý Đấu trường Trực tuyến")
    add_body_p("Người chơi kết nối Máy chủ → Chọn Tạo phòng Đấu trường (hoặc Nhập mã 6 số) → Thêm AI Bot nếu thiếu người → Bấm Sẵn sàng (Ready) → Đếm ngược 30s thi đấu Click realtime → Server tổng kết Bảng xếp hạng trao thưởng.")

    add_h2("3.3. Sơ đồ Luồng xử lý Hợp tác Trực tuyến")
    add_body_p("Người chơi vào Phòng Hợp tác → Cùng các thành viên nhấp chuột thu gom nguyên liệu Thịt/Gỗ/Đá → Đóng góp nguyên liệu mua Nâng cấp chung → Đánh gục Boss/Cổ thụ/Vỉa đá → Hệ thống lưu vết nhật ký đồng bộ.")

    # ==========================================
    # CHƯƠNG 4
    # ==========================================
    add_h1("CHƯƠNG 4: THIẾT KẾ CƠ SỞ DỮ LIỆU VÀ MÁY CHỦ MULTIPLAYER")

    add_h2("4.1. Cấu trúc Thực thể Cơ sở Dữ liệu (Database Schema)")
    add_body_p("Cơ sở dữ liệu MySQL của hệ thống được thiết kế tối ưu với các bảng lưu trữ thông tin phòng chơi, tài khoản và lịch sử trận đấu như sau:")

    # Table 1: Schema Table
    table1 = doc.add_table(rows=6, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table1.rows[0].cells
    headers = ['Tên Bảng', 'Khóa Chính (PK)', 'Khóa Ngoại (FK)', 'Mô Tả Chức Năng']
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_bg(hdr_cells[idx], "1E293B") # Dark Charcoal header
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    data1 = [
        ['players', 'id (VARCHAR)', 'room_code', 'Quản lý thông tin người chơi, biệt danh, chỉ số DPC, DPS, điểm số.'],
        ['rooms', 'code (VARCHAR)', '-', 'Quản lý thông tin phòng chơi online, trạng thái (lobby/playing/finished), timer.'],
        ['game_sessions', 'id (INT AUTO)', 'room_code', 'Lưu vết lịch sử các trận đấu đã hoàn thành và người chiến thắng.'],
        ['coop_resources', 'id (INT AUTO)', 'room_code', 'Lưu giữ số lượng tài nguyên Thịt, Gỗ, Đá tích lũy trong chế độ Hợp tác.'],
        ['coop_upgrades', 'id (INT AUTO)', 'room_code', 'Lưu giữ cấp độ các nâng cấp dùng chung của phòng Hợp tác.']
    ]

    for row_idx, row_data in enumerate(data1, start=1):
        row_cells = table1.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_bg(row_cells[col_idx], bg_color)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx != 1 else WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)

    add_body_p("") # spacing

    add_h2("4.2. Xử lý các Tình huống Vận hành Thực tế")
    add_bullet("Nếu thiết bị đột ngột mất kết nối Socket, game tự động chuyển về chế độ Offline mà không làm gián đoạn trải nghiệm người chơi.", "Tình huống 1: Mất kết nối mạng đột ngột (Network Disconnect Fallback):")
    add_bullet("Chủ phòng có thể bấm Thêm Bot. Hệ thống tự động khởi tạo luồng AI Bot tự động nhấp chuột với tần số ngẫu nhiên 200-400ms.", "Tình huống 2: Thiếu người chơi trong phòng Đấu trường (AI Bot Filling):")
    add_bullet("Khi deploy lên máy chủ Railway Cloud, server tự động ràng buộc địa chỉ `0.0.0.0:${PORT}` để tiếp nhận kết nối bên ngoài mượt mà.", "Tình huống 3: Tương thích máy chủ Cloud Railway (Host Binding):")

    # ==========================================
    # CHƯƠNG 5
    # ==========================================
    add_h1("CHƯƠNG 5: THIẾT LẬP CÔNG NGHỆ VÀ TRIỂN KHAI CLOUD")

    # Tech Table
    table2 = doc.add_table(rows=5, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Phân Loại Công Nghệ"
    hdr2[1].text = "Danh Sách Công Nghệ & Thư Viện Sử Dụng"
    for cell in hdr2:
        set_cell_bg(cell, "1E293B")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    data2 = [
        ['Frontend UI Engine', 'React 19, HTML5, Vanilla CSS3 (Glassmorphism & Sunburst animations), Lucide Icons, Web Audio API'],
        ['Backend Core Server', 'Node.js, Express.js (REST API & Static Server)'],
        ['Real-time Protocol', 'Socket.io v4 (Full-duplex WebSocket Real-time Communication)'],
        ['Database & Deployment', 'MySQL (mysql2 pool), Railway Cloud (Node Server), GitHub Pages (Static Host), GitHub Actions CI/CD']
    ]

    for r_idx, r_data in enumerate(data2, start=1):
        row_cells = table2.rows[r_idx].cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            set_cell_bg(row_cells[c_idx], bg_color)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if c_idx == 0:
                    run.font.bold = True

    add_body_p("")
    add_h1("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    add_body_p("Đồ án Game Siêu Clicker Tam Hợp đã hoàn thành đúng tiến độ, đáp ứng đầy đủ các yêu cầu kỹ thuật và tính năng đặt ra. Hệ thống không chỉ mang lại trải nghiệm giải trí ấn tượng với đồ họa hoạt hình 3D mượt mà mà còn khẳng định tính khả thi của việc kết hợp các công nghệ Web hiện đại (React 19, Node.js, Socket.io, MySQL) trong các bài toán thực tế.")
    add_body_p("Trong tương lai, nhóm sẽ tiếp tục phát triển thêm các tính năng như: Bảng xếp hạng toàn cầu (Global Leaderboard), tích hợp ví vật phẩm NFT và mở rộng thêm nhiều đồ họa Theme phong phú hơn nữa.")

    output_path = 'BaoCao_GameClicker.docx'
    doc.save(output_path)
    print(f"Perfect game report saved to {output_path} successfully!")

if __name__ == '__main__':
    build_perfect_game_report()
