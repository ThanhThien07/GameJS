import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_element(name):
    return docx.oxml.OxmlElement(name)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = create_element('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = create_element(f'w:{m}')
        node.set(docx.oxml.ns.qn('w:w'), str(val))
        node.set(docx.oxml.ns.qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_report():
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base colors
    primary_color = RGBColor(121, 40, 202) # Purple #7928ca
    secondary_color = RGBColor(255, 0, 128) # Pink #ff0080
    dark_gray = RGBColor(40, 40, 40)
    
    # Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO ĐỒ ÁN MÔN HỌC")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = primary_color
    title_p.paragraph_format.space_before = Pt(40)
    title_p.paragraph_format.space_after = Pt(10)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("PHÁT TRIỂN ỨNG DỤNG WEB GAME TAP TAP CLICKER MULTIPLAYER\n(Chế độ Offline & Online Đấu Trường và Hợp Tác)")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(16)
    sub_run.font.bold = True
    sub_run.font.color.rgb = secondary_color
    sub_p.paragraph_format.space_after = Pt(50)

    info_data = [
        ("Sinh viên thực hiện:", "Nguyễn Hoàng Hùng"),
        ("Mã số sinh viên:", "501250384"),
        ("Đề tài dự án:", "Web Game Tap Tap Clicker Multiplayer"),
        ("GitHub Repository:", "https://github.com/ThanhThien07/GameJS.git"),
        ("Link Web GitHub Pages:", "https://thanhthien07.github.io/GameJS/")
    ]
    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(info_data):
        row = info_table.rows[i]
        
        run_l = row.cells[0].paragraphs[0].add_run(label)
        run_l.font.name = 'Arial'
        run_l.font.size = Pt(11)
        run_l.font.bold = True
        
        run_v = row.cells[1].paragraphs[0].add_run(val)
        run_v.font.name = 'Arial'
        run_v.font.size = Pt(11)
        run_v.font.color.rgb = dark_gray
        
        set_cell_margins(row.cells[0], top=100, bottom=100, left=150, right=150)
        set_cell_margins(row.cells[1], top=100, bottom=100, left=150, right=150)

    doc.add_page_break()

    # SECTION 1: ĐẶT VẤN ĐỀ
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Đặt Vấn Đề: Lý Do Chọn Đề Tài")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = primary_color
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after = Pt(10)

    p1 = doc.add_paragraph()
    p1_run = p1.add_run(
        "Trò chơi điện tử dạng \"Tap Tap\" hay \"Clicker\" là một trong những thể loại game đơn giản nhất nhưng có sức hút "
        "kỳ lạ và khả năng giữ chân người chơi cực kỳ tốt. Người chơi chỉ cần thực hiện thao tác nhấp chuột liên tiếp vào "
        "màn hình để kiếm tài nguyên, nâng cấp sức mạnh, mua các cỗ máy tự động để gia tăng tốc độ tích lũy tiền tệ vô hạn. "
        "Tuy nhiên, hầu hết các tựa game clicker hiện nay trên thị trường là game chơi đơn (Single-player), người chơi dễ cảm "
        "thấy nhàm chán sau một thời gian dài lặp đi lặp lại hành động một mình."
    )
    p1_run.font.name = 'Arial'
    p1_run.font.size = Pt(11)
    p1.paragraph_format.space_after = Pt(10)

    p2 = doc.add_paragraph()
    p2_run = p2.add_run(
        "Nhận thấy cơ hội nâng tầm thể loại này, sinh viên Nguyễn Hoàng Hùng (MSSV: 501250384) đã quyết định thực hiện đề tài \"Phát triển ứng dụng Web Game "
        "Tap Tap Clicker Multiplayer\". Dự án tập trung giải quyết các bài toán:"
    )
    p2_run.font.name = 'Arial'
    p2_run.font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(8)

    bullet_points = [
        ("Trải nghiệm kết nối cộng đồng: ", "Hỗ trợ 3 người chơi cùng lúc trong phòng đấu trường Competitive để xem ai kiếm được nhiều tiền hơn trong thời gian giới hạn 60 giây, hoặc cùng tham gia phòng đặc biệt Co-op (Chế độ 3) để nhặt gỗ, đào đá, săn thịt nâng cấp các chỉ số sức mạnh chung."),
        ("Linh hoạt Online/Offline: ", "Trò chơi tự động nhận dạng trạng thái kết nối mạng của thiết bị. Khi có mạng, người chơi thoải mái tạo phòng thi đấu thời gian thực. Khi mất mạng, game tự động chuyển về chế độ offline và vẫn cho phép chơi bình thường, bảo toàn dữ liệu nâng cấp lưu tại Local Storage."),
        ("Đa dạng mô hình tương tác: ", "Tích hợp 3 chủ đề đồ họa sinh động bao gồm Đánh Quái Vật (Monster Fight), Chặt Gỗ (Woodcutting), và Đào Đá (Stone Mining) với hoạt ảnh tương tác riêng cho từng chủ đề, tránh sự nhàm chán thị giác.")
    ]

    for label, desc in bullet_points:
        bp = doc.add_paragraph(style='List Bullet')
        run_label = bp.add_run(label)
        run_label.font.name = 'Arial'
        run_label.font.size = Pt(11)
        run_label.font.bold = True
        run_desc = bp.add_run(desc)
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(11)
        bp.paragraph_format.space_after = Pt(4)

    # SECTION 2: GIẢI QUYẾT VẤN ĐỀ
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Giải Quyết Vấn Đề: Công Nghệ & Kiến Trúc F+B")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(18)
    h2_run.font.bold = True
    h2_run.font.color.rgb = primary_color
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after = Pt(10)

    p3 = doc.add_paragraph()
    p3_run = p3.add_run(
        "Để giải quyết các yêu cầu đặt ra về mặt kỹ thuật, hệ thống được thiết kế theo mô hình Full-stack với kiến trúc tách biệt "
        "Frontend và Backend (F+B), giao tiếp thời gian thực thông qua giao thức WebSockets."
    )
    p3_run.font.name = 'Arial'
    p3_run.font.size = Pt(11)
    p3.paragraph_format.space_after = Pt(10)

    # Sub-heading 2.1
    h2_1 = doc.add_paragraph()
    h2_1_run = h2_1.add_run("2.1 Công nghệ sử dụng ở phía Frontend (F)")
    h2_1_run.font.name = 'Arial'
    h2_1_run.font.size = Pt(13)
    h2_1_run.font.bold = True
    h2_1_run.font.color.rgb = secondary_color
    h2_1.paragraph_format.space_before = Pt(10)
    h2_1.paragraph_format.space_after = Pt(6)

    p4 = doc.add_paragraph()
    p4_run = p4.add_run(
        "Frontend được xây dựng dựa trên React.js kết hợp với trình đóng gói siêu tốc Vite. Lựa chọn này giúp giao diện "
        "render cực kỳ mượt mà nhờ cơ chế Virtual DOM của React, điều vô cùng quan trọng đối với game clicker có tần suất "
        "nhấp chuột và thay đổi điểm số liên tục. Chi tiết bao gồm:"
    )
    p4_run.font.name = 'Arial'
    p4_run.font.size = Pt(11)
    p4.paragraph_format.space_after = Pt(8)

    fe_bullets = [
        ("React hooks (useState, useEffect, useRef): ", "Quản lý luồng dữ liệu, trạng thái nâng cấp, thời gian bộ đếm ngược, và liên kết cổng kết nối socket."),
        ("Bố cục 3 cột chuẩn Mine Clicker (3-Column Layout): ", "Giao diện thiết kế theo bố cục 3 cột chuẩn mực của game Mine Clicker trên CrazyGames: Cột trái nâng cấp công cụ thủ công (Click Tools), cột giữa chứa đối tượng nhấp chuột chính và thanh HP/thanh nộ, cột phải là tuyển nhân công tự động (Auto Helpers). Toàn bộ sử dụng phong cách sáng tối giản, nền trắng trang nhã kết hợp các hàng nâng cấp dạng thanh ngang trực quan."),
        ("Socket.io Client: ", "Kết nối thời gian thực hai chiều, gửi tín hiệu click chuột lên server và nhận về danh sách xếp hạng tức thời."),
        ("Hệ thống nộ x2 & Thanh Combo Streak: ", "Khi người chơi click liên tiếp dưới 450ms, thanh năng lượng sẽ sạc đầy kích hoạt nộ x2 trong 6s và tích lũy bộ nhân Combo Streak lên đến x3 tiền thưởng."),
        ("Hệ thống Âm thanh Web Audio API (`src/utils/audio.js`): ", "Tự động phát sinh hiệu ứng âm thanh click, mua đồ, dùng kỹ năng và trùng sinh mượt mà bằng sóng âm tổng hợp Web Audio API (không cần tải file MP3 ngoài)."),
        ("Kỹ Năng Nguồn Lực Chủ Động (Active Skills): ", "Tích hợp kỹ năng 'Cơn Cuồng Phong' (Nhân đôi DPC trong 10s) và 'Bão Vàng' (Thưởng vàng tức thì dựa trên tốc độ đào hiện tại)."),
        ("Hệ thống Trùng Sinh (Prestige / Rebirth): ", "Cho phép người chơi reset điểm số và cấp độ nâng cấp thường để đổi lấy Tinh Thể Linh Hồn (Soul Crystals 💎) giúp tăng vĩnh viễn +15% sức mạnh cho tất cả các lần chơi tiếp theo."),
        ("Bảng Thành Tựu (Achievements Modal): ", "Bảng danh hiệu và mốc thưởng giúp tăng chiều sâu chiến thuật và giữ chân người chơi lâu dài.")
    ]
    for label, desc in fe_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        run_label = bp.add_run(label)
        run_label.font.name = 'Arial'
        run_label.font.size = Pt(11)
        run_label.font.bold = True
        run_desc = bp.add_run(desc)
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(11)
        bp.paragraph_format.space_after = Pt(4)

    # Sub-heading 2.2
    h2_2 = doc.add_paragraph()
    h2_2_run = h2_2.add_run("2.2 Công nghệ sử dụng ở phía Backend (B)")
    h2_2_run.font.name = 'Arial'
    h2_2_run.font.size = Pt(13)
    h2_2_run.font.bold = True
    h2_2_run.font.color.rgb = secondary_color
    h2_2.paragraph_format.space_before = Pt(10)
    h2_2.paragraph_format.space_after = Pt(6)

    p5 = doc.add_paragraph()
    p5_run = p5.add_run(
        "Backend sử dụng nền tảng Node.js với Express framework và Socket.io để xử lý logic đa người chơi tập trung trên server, "
        "đảm bảo tính nhất quán dữ liệu giữa 3 người chơi trong phòng. Chi tiết:"
    )
    p5_run.font.name = 'Arial'
    p5_run.font.size = Pt(11)
    p5.paragraph_format.space_after = Pt(8)

    be_bullets = [
        ("Quản lý phòng (Room Manager): ", "Lưu trữ động trạng thái các phòng game trong bộ nhớ RAM, phân định rõ phòng Thi Đấu (Competitive) và phòng Hợp Tác (Co-op). Khởi tạo mã phòng ngẫu nhiên 6 ký tự để người chơi tìm kiếm và kết nối."),
        ("Đồng bộ thời gian thực (Socket.io Server): ", "Lắng nghe sự kiện click từ các client, tính toán sức mạnh cộng dồn và phát (broadcast) bảng xếp hạng mới nhất cho tất cả các thành viên trong phòng. Đối với chế độ Co-op, đồng bộ kho tài nguyên chung và cấp độ nâng cấp công nghệ chung."),
        ("Hệ thống AI Bots (Người chơi ảo): ", "Để giải quyết bài toán thiếu người chơi khi test hoặc chơi một mình, Server tích hợp tính năng tạo Bot tự động. Khi được kích hoạt, Bot sẽ tự động gia nhập phòng, đánh dấu sẵn sàng, thực hiện nhấp chuột giả lập và tự động dùng tiền kiếm được mua nâng cấp tối ưu hóa điểm số để cạnh tranh sòng phẳng với người chơi thực."),
        ("Đồng bộ bộ đếm ngược (Game Timer): ", "Bộ đếm thời gian 60 giây chạy trực tiếp trên Server và gửi nhịp tích tắc xuống client để chống gian lận thời gian ở máy khách.")
    ]
    for label, desc in be_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        run_label = bp.add_run(label)
        run_label.font.name = 'Arial'
        run_label.font.size = Pt(11)
        run_label.font.bold = True
        run_desc = bp.add_run(desc)
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(11)
        bp.paragraph_format.space_after = Pt(4)

    # Sub-heading 2.3
    h2_3 = doc.add_paragraph()
    h2_3_run = h2_3.add_run("2.3 Quy trình Triển khai Tự động (CI/CD & GitHub Pages)")
    h2_3_run.font.name = 'Arial'
    h2_3_run.font.size = Pt(13)
    h2_3_run.font.bold = True
    h2_3_run.font.color.rgb = secondary_color
    h2_3.paragraph_format.space_before = Pt(10)
    h2_3.paragraph_format.space_after = Pt(6)

    p_cicd = doc.add_paragraph()
    p_cicd_run = p_cicd.add_run(
        "Dự án được tự động hóa quy trình đóng gói và triển khai (CI/CD) lên nền tảng GitHub Pages thông qua GitHub Actions "
        "(`.github/workflows/deploy.yml`). Khi có bất kỳ thay đổi nào được push lên branch main, hệ thống sẽ tự động khởi tạo môi trường "
        "Node.js, đóng gói bản build Vite và deploy giao diện web lên địa chỉ công khai:"
    )
    p_cicd_run.font.name = 'Arial'
    p_cicd_run.font.size = Pt(11)
    p_cicd.paragraph_format.space_after = Pt(8)

    cicd_bullets = [
        ("Cấu hình Vite Base Path (`vite.config.js`): ", "Thiết lập base: '/GameJS/' để đảm bảo các tài nguyên tĩnh (JavaScript, CSS, Assets) tải chính xác trên GitHub Pages."),
        ("Tự động hóa build với GitHub Actions: ", "Workflow tự động thực hiện các lệnh `npm ci`, `npm run build` và đẩy thư mục `dist/` lên môi trường gh-pages trong chưa đầy 1 phút."),
        ("Địa chỉ truy cập trực tuyến: ", "https://thanhthien07.github.io/GameJS/")
    ]
    for label, desc in cicd_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        run_label = bp.add_run(label)
        run_label.font.name = 'Arial'
        run_label.font.size = Pt(11)
        run_label.font.bold = True
        run_desc = bp.add_run(desc)
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(11)
        bp.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # SECTION 3: KẾT THÚC VẤN ĐỀ
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Kết Thúc Vấn Đề: Đánh Giá & Định Hướng Phát Triển")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(18)
    h3_run.font.bold = True
    h3_run.font.color.rgb = primary_color
    h3.paragraph_format.space_before = Pt(20)
    h3.paragraph_format.space_after = Pt(10)

    p6 = doc.add_paragraph()
    p6_run = p6.add_run(
        "Sau thời gian nghiên cứu và thực hiện, dự án đã cơ bản hoàn thành đầy đủ tất cả các yêu cầu đề ra của một tựa game clicker hiện đại."
    )
    p6_run.font.name = 'Arial'
    p6_run.font.size = Pt(11)
    p6.paragraph_format.space_after = Pt(12)

    # Analysis table for Pros/Cons
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("ƯU ĐIỂM CỦA DỰ ÁN").font.bold = True
    hdr_cells[1].paragraphs[0].add_run("HẠN CHẾ & KHUYẾT ĐIỂM").font.bold = True
    
    for cell in hdr_cells:
        set_cell_background(cell, "7928ca")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

    # Data Rows
    row1_cells = table.rows[1].cells
    row1_cells[0].paragraphs[0].add_run(
        "- Bố cục 3 cột chuẩn game Mine Clicker (CrazyGames) trực quan, màu nền sáng nhẹ trang nhã.\n"
        "- Khả năng chuyển đổi online/offline linh hoạt, nhận diện sự cố mạng thông minh.\n"
        "- Hệ thống Multiplayer thời gian thực qua socket hoạt động ổn định, không giật lag.\n"
        "- Tích hợp hiệu ứng âm thanh sinh động bằng Web Audio API (không cần tải mp3 ngoại vi)."
    ).font.name = 'Arial'
    row1_cells[1].paragraphs[0].add_run(
        "- Chưa tích hợp cơ sở dữ liệu (Database) để lưu trữ vĩnh viễn cấp độ nâng cấp offline hay tài khoản của người dùng khi khởi động lại server."
    ).font.name = 'Arial'

    row2_cells = table.rows[2].cells
    row2_cells[0].paragraphs[0].add_run(
        "- Tích hợp AI Bots giúp trải nghiệm multiplayer có thể kiểm thử dễ dàng bởi 1 người chơi duy nhất.\n"
        "- Chế độ Hợp tác (Chế độ 3) có cơ chế nhặt đồ nâng cấp độc đáo, thúc đẩy tương tác đồng đội.\n"
        "- Hệ thống Trùng Sinh (Prestige), Kỹ Năng Nguồn Lực và Thành Tựu tăng chiều sâu và độ cuốn hút dài hạn."
    ).font.name = 'Arial'
    row2_cells[1].paragraphs[0].add_run(
        "- Chưa có cơ chế bảo mật chống các phần mềm bên thứ ba tự động nhấp chuột (Auto-click hack) để gian lận điểm số."
    ).font.name = 'Arial'

    for row in table.rows[1:]:
        for idx, cell in enumerate(row.cells):
            set_cell_background(cell, "f3e8ff" if idx == 0 else "fff1f2")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.color.rgb = dark_gray

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Future direction
    h3_1 = doc.add_paragraph()
    h3_1_run = h3_1.add_run("3.1 Hướng phát triển tương lai")
    h3_1_run.font.name = 'Arial'
    h3_1_run.font.size = Pt(13)
    h3_1_run.font.bold = True
    h3_1_run.font.color.rgb = secondary_color
    h3_1.paragraph_format.space_before = Pt(10)
    h3_1.paragraph_format.space_after = Pt(6)

    p7 = doc.add_paragraph()
    p7_run = p7.add_run(
        "Trong các phiên bản tiếp theo, dự án dự kiến sẽ bổ sung cơ sở dữ liệu MongoDB để lưu trữ bảng xếp hạng người chơi toàn quốc, "
        "kết nối API âm thanh Web Audio API sinh động, xây dựng hệ thống bang hội nâng cấp kỹ năng chiến đấu dài hạn và áp dụng thuật toán "
        "kiểm tra độ trễ click chuột để chặn đứng các công cụ gian lận autoclick."
    )
    p7_run.font.name = 'Arial'
    p7_run.font.size = Pt(11)
    p7.paragraph_format.space_after = Pt(15)

    # SECTION 4: PHỤ LỤC MÃ NGUỒN
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Phụ Lục Mã Nguồn Đặc Trưng")
    h4_run.font.name = 'Arial'
    h4_run.font.size = Pt(18)
    h4_run.font.bold = True
    h4_run.font.color.rgb = primary_color
    h4.paragraph_format.space_before = Pt(20)
    h4.paragraph_format.space_after = Pt(10)

    p8 = doc.add_paragraph()
    p8_run = p8.add_run("Dưới đây là một số đoạn mã nguồn cốt lõi giải quyết bài toán kết nối phòng thời gian thực và tự động tạo bot ở Server:")
    p8_run.font.name = 'Arial'
    p8_run.font.size = Pt(11)
    p8.paragraph_format.space_after = Pt(8)

    # Code block container
    code_text = (
        "// Đoạn mã tạo AI Bot trên server.js để đồng bộ phòng 3 người\n"
        "socket.on('addBot', ({ code }, callback) => {\n"
        "  const room = rooms.get(code);\n"
        "  if (!room) return callback({ success: false, message: 'Không tìm thấy phòng!' });\n"
        "  if (room.players.length >= 3) return callback({ success: false, message: 'Phòng đã đầy!' });\n"
        "  \n"
        "  const botNames = ['SuperClicker_AI', 'IronFinger_Bot', 'TappingMaster_AI', 'SwiftClick_AI'];\n"
        "  const randomName = botNames[Math.floor(Math.random() * botNames.length)];\n"
        "  \n"
        "  const botPlayer = {\n"
        "    id: `bot_${Math.random()}`,\n"
        "    name: `${randomName} (AI)`,\n"
        "    score: 0,\n"
        "    clicks: 0,\n"
        "    dpc: 1,\n"
        "    dps: 0,\n"
        "    upgrades: { clicker: 0, pickaxe: 0, minecart: 0 },\n"
        "    isBot: true,\n"
        "    isReady: true // AI luôn sẵn sàng đấu\n"
        "  };\n"
        "  \n"
        "  room.players.push(botPlayer);\n"
        "  io.to(code).emit('roomUpdated', room);\n"
        "  callback({ success: true, room });\n"
        "});"
    )

    code_p = doc.add_paragraph()
    code_run = code_p.add_run(code_text)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9.5)
    code_run.font.color.rgb = RGBColor(50, 50, 50)
    
    # Adding a light grey shading to the paragraph for coding blocks
    pBrd = parse_xml(f'<w:pBrd {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="7928ca"/></w:pBrd>')
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f9f5ff"/>')
    code_p._p.get_or_add_pPr().append(pBrd)
    code_p._p.get_or_add_pPr().append(shd)

    # Save document
    filename = "BaoCao_GameClicker.docx"
    doc.save(filename)
    print(f"Word report successfully generated at: {os.path.abspath(filename)}")

if __name__ == "__main__":
    generate_report()
