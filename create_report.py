import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = docx.Document()

    # Page Margins: 2.54cm (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles helper
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    def set_cell_bg(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def add_p(text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, size=13, color=(30, 41, 59), space_before=3, space_after=6, line_spacing=1.15):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_h1(text):
        p = add_p(text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=16, color=(124, 58, 237), space_before=16, space_after=8)
        return p

    def add_h2(text):
        p = add_p(text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=14, color=(15, 23, 42), space_before=12, space_after=6)
        return p

    def add_h3(text):
        p = add_p(text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, italic=True, size=13, color=(71, 85, 105), space_before=8, space_after=4)
        return p

    def add_bullet(text, bold_prefix="", level=0):
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        
        run_bullet = p.add_run("•  ")
        run_bullet.font.name = 'Times New Roman'
        run_bullet.font.size = Pt(13)
        run_bullet.font.color.rgb = RGBColor(124, 58, 237)

        if bold_prefix:
            run_b = p.add_run(bold_prefix + " ")
            run_b.font.name = 'Times New Roman'
            run_b.font.size = Pt(13)
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(15, 23, 42)

        run_t = p.add_run(text)
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(13)
        run_t.font.color.rgb = RGBColor(30, 41, 59)
        return p

    # ==========================================
    # 1. TRANG BÌA (COVER PAGE)
    # ==========================================
    add_p("BỘ GIÁO DỤC VÀ ĐÀO TẠO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, color=(15, 23, 42), space_before=0, space_after=2)
    add_p("TRƯỜNG CAO ĐẲNG CÔNG NGHỆ THÔNG TIN TP. HỒ CHÍ MINH", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, color=(15, 23, 42), space_before=0, space_after=2)
    add_p("KHOA CÔNG NGHỆ THÔNG TIN – ĐIỆN TỬ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, color=(71, 85, 105), space_before=0, space_after=24)

    add_p("PHÂN TÍCH THIẾT KẾ VÀ XÂY DỰNG GAME SIÊU CLICKER TAM HỢP", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20, color=(124, 58, 237), space_before=12, space_after=4)
    add_p("(TAP TAP CLICKER MULTIPLAYER)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, color=(219, 39, 119), space_before=0, space_after=12)
    add_p("HỌC KỲ 3, NĂM HỌC 2025 - 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, color=(100, 116, 139), space_before=0, space_after=40)

    # Info Block Box
    add_p("NHÓM THỰC HIỆN: NGUYỄN HOÀNG HÙNG & HOÀNG THỊ DIỄM LY", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=13, color=(15, 23, 42), space_before=4, space_after=6)
    add_p("NSVTH:   1. NGUYỄN HOÀNG HÙNG  (MSSV: 501250384)", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=13, color=(30, 41, 59), space_before=2, space_after=3)
    add_p("               2. HOÀNG THỊ DIỄM LY", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=13, color=(30, 41, 59), space_before=0, space_after=6)
    add_p("MÔN :      JAVASCRIPTS", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=13, color=(15, 23, 42), space_before=2, space_after=4)
    add_p("LỚP:        CD25CT6", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=13, color=(15, 23, 42), space_before=2, space_after=4)
    add_p("GVHD:     ThS. Trương Châu Long", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=13, color=(124, 58, 237), space_before=2, space_after=40)

    add_p("Tp. Hồ Chí Minh, tháng 07 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, color=(100, 116, 139), space_before=30, space_after=0)

    doc.add_page_break()

    # ==========================================
    # 2. LỜI CẢM ƠN
    # ==========================================
    add_h1("LỜI CẢM ƠN")
    add_p("Lời đầu tiên, nhóm chúng em xin gửi lời cảm ơn chân thành đến Ban Giám hiệu Trường Cao đẳng Công nghệ Thông tin TP. Hồ Chí Minh cùng toàn thể quý Thầy Cô Khoa Công nghệ Thông tin – Điện tử đã tạo điều kiện học tập tốt nhất, giảng dạy những nền tảng kiến thức chuyên môn vững chắc và truyền cảm hứng sáng tạo cho chúng em trong suốt quá trình học tập.")
    add_p("Đặc biệt, nhóm em xin bày tỏ lòng biết ơn sâu sắc nhất đến thầy ThS. Trương Châu Long — giảng viên phụ trách môn JAVASCRIPTS. Thầy đã tận tình hướng dẫn, truyền đạt kiến thức lập trình hiện đại, chỉ dẫn cách tối ưu hóa hiệu năng ứng dụng Web thời gian thực và luôn đưa ra những lời khuyên, định hướng quý báu giúp nhóm hoàn thiện đồ án này một cách bài bản nhất.")
    add_p("Mặc dù nhóm đã đầu tư nghiên cứu và hoàn thiện sản phẩm với sự quyết tâm cao, song do kinh nghiệm thực tế còn hạn chế nên khó tránh khỏi những thiếu sót nhất định. Nhóm rất mong nhận được những góp ý và đánh giá chân thành từ thầy ThS. Trương Châu Long để sản phẩm ngày càng hoàn thiện hơn.")
    add_p("Nhóm xin trân trọng và chân thành cảm ơn!")

    # ==========================================
    # 3. MỞ ĐẦU
    # ==========================================
    add_h1("MỞ ĐẦU")
    add_p("Trong kỷ nguyên số hóa hiện nay, ngành công nghiệp giải trí kỹ thuật số và phát triển Game Web trên nền tảng JavaScript đang bùng nổ mạnh mẽ nhờ tính linh hoạt, khả năng tương thích cao trên mọi thiết bị di động cũng như máy tính mà không cần cài đặt phức tạp. Dòng game Clicker (Idle Clicker Game) từ lâu đã trở thành một hiện tượng thu hút hàng triệu người chơi nhờ lối chơi đơn giản nhưng cực kỳ cuốn hút, giúp giải tỏa căng thẳng nhanh chóng.")
    add_p("Tuy nhiên, phần lớn các tựa game Clicker truyền thống hiện nay chỉ dừng lại ở chế độ chơi đơn (Offline) thụ động, thiếu tính kết nối cộng đồng và dễ gây cảm giác nhàm chán sau một thời gian ngắn. Xuất phát từ nhu cầu thực tiễn đó, nhóm em đã lựa chọn đề tài \"Phân tích thiết kế và xây dựng Game Siêu Clicker Tam Hợp (Tap Tap Clicker Multiplayer)\" để tạo ra một sản phẩm game Web hiện đại, kết hợp hài hòa giữa tính năng chơi đơn tích lũy và chế độ thi đấu/hợp tác trực tuyến đa người chơi (Real-time Multiplayer) thông qua công nghệ WebSocket tiên tiến.")

    # ==========================================
    # 4. GIỚI THIỆU PHẦN MỀM GAME CLICKER
    # ==========================================
    add_h1("1. GIỚI THIỆU PHẦN MỀM GAME CLICKER")
    
    add_h2("1.1. Lý do chọn đề tài")
    add_bullet("Các tựa game Web truyền thống thường thiếu sự tương tác thời gian thực giữa các người chơi với nhau.", "Hạn chế tương tác:")
    add_bullet("Nhiều sản phẩm game Clicker hiện tại có giao diện đơn điệu, đồ họa thiếu sức sống và thiếu cơ chế thưởng phong phú.", "Đồ họa và lối chơi:")
    add_bullet("Cần áp dụng các công nghệ JavaScript hiện đại (React 19, Vite, Node.js, Socket.io, MySQL) vào sản phẩm thực tế có khả năng đưa lên môi trường Cloud (Railway, GitHub Pages).", "Ứng dụng công nghệ mới:")

    add_h2("1.2. Lợi ích mang lại của phần mềm Game")
    add_bullet("Trải nghiệm game mượt mà trên trình duyệt Web không cần cài đặt.", "Trải nghiệm đa nền tảng:")
    add_bullet("Chơi mọi lúc mọi nơi kể cả khi không có kết nối mạng (Tự động lưu tiến trình vào LocalStorage).", "Chế độ Offline linh hoạt:")
    add_bullet("Kết nối người chơi trực tuyến trong đấu trường 1v1v1 hoặc phòng Hợp tác 3 người với AI Bot hỗ trợ thông minh.", "Tính năng Multiplayer real-time:")
    add_bullet("Đồ họa hoạt hình 3D bắt mắt với 3 mô hình tương tác độc đáo (Đánh Quái Vật ⚔️, Chặt Gỗ 🪵, Đào Đá Quặng 🪨).", "Đa dạng đồ họa Theme:")

    # ==========================================
    # 5. TỔNG QUAN PHẦN MỀM GAME
    # ==========================================
    add_h1("2. TỔNG QUAN PHẦN MỀM GAME")
    add_p("Hệ thống phần mềm Siêu Clicker Tam Hợp được thiết kế với kiến trúc mô-đun hóa, chia làm các phân hệ quản lý chức năng chính như sau:")

    add_h2("2.1. Quản lý Chế độ Chơi Ngoại tuyến (Offline Mode)")
    add_bullet("Cho phép người chơi nhấp chuột thủ công để tạo ra vàng/sát thương. Tích lũy điểm để tăng chỉ số DPC (Damage Per Click).", "Tương tác nhấp chuột (Manual Click):")
    add_bullet("Thuê các công cụ và cỗ máy tự động (Găng tay sắt, Rìu chặt củi, Kiếm kim cương, Xe goòng mỏ, Máy khoan laze, Giàn khoan siêu cấp) để tạo nguồn thu nhập tự động DPS (Damage Per Second).", "Hệ thống Nâng cấp tự động (Auto Helpers):")
    add_bullet("Khi đạt 50,000 vàng, người chơi có thể Trùng Sinh để quy đổi lấy Tinh Thể Linh Hồn (Soul Crystals), tăng vĩnh viễn +15% tất cả chỉ số.", "Hệ thống Trùng Sinh (Prestige/Rebirth):")
    add_bullet("Kích hoạt Cơn Cuồng Phong (x2 DPC trong 10s) và Bão Vàng (Thưởng vàng tức thì).", "Bộ Kỹ năng Nguồn lực (Active Skills):")
    add_bullet("Theo dõi và mở khóa các mốc thành tựu (100 Clicks, 100k Vàng, Bậc thầy Trùng sinh).", "Bảng Thành tựu (Achievements):")

    add_h2("2.2. Quản lý Chế độ Trực tuyến Đấu trường (Competitive Online Mode)")
    add_bullet("Hệ thống khởi tạo mã phòng 6 ký tự ngẫu nhiên (VD: X9A2B1) cho phép người chơi tạo hoặc gia nhập phòng.", "Tạo & Quản lý Phòng đấu (Room System):")
    add_bullet("Cho phép chủ phòng thêm các đối thủ máy thông minh (AI Bots) để lấp đầy suất chơi khi thiếu người.", "Tích hợp AI Bot thông minh:")
    add_bullet("Đếm ngược 30 giây thi đấu kịch tính. Người chơi cạnh tranh điểm số realtime qua kết nối WebSocket.", "Đấu trường thời gian thực (Real-time Match):")
    add_bullet("Kết xuất Bảng vàng xếp hạng thứ hạng #1, #2, #3 và lưu vết kết quả đấu trường.", "Bảng xếp hạng Trận đấu:")

    add_h2("2.3. Quản lý Chế độ Trực tuyến Hợp tác (Co-op Online Mode)")
    add_bullet("3 người chơi cùng nhấp chuột hạ gục Mục tiêu chung.", "Hợp lực đánh Boss/Khai mỏ:")
    add_bullet("Các tài nguyên Thịt 🥩, Gỗ 🪵, Đá 🪨 ngẫu nhiên rơi ra được đồng bộ ngay lập tức tới tất cả thành viên trong phòng.", "Rơi tài nguyên ngẫu nhiên (Resource Drops):")
    add_bullet("Người chơi đóng góp tài nguyên thu gom được để mua nâng cấp Sát thương chung, Hệ số nhân và Robot Auto-Click cho toàn phòng.", "Nâng cấp dùng chung (Shared Upgrades):")
    add_bullet("Hiển thị dòng tin nhắn thông báo mỗi khi có thành viên nhặt tài nguyên hoặc kích hoạt nâng cấp.", "Nhật ký hành động (Action Logs):")

    add_h2("2.4. Quản lý Đồ họa & Hiệu ứng Âm thanh (Themes & Audio Synthesizer)")
    add_bullet("Đánh Quái Vật (Monster Fight), Tiều Phu Chặt Gỗ (Woodcutter), Thợ Mỏ Đào Đá (Stone Mining).", "3 Đồ họa Theme 3D Hoạt hình:")
    add_bullet("Sử dụng Web Audio API tổng hợp âm thanh trực tiếp bằng mã lệnh (Click, Mua hàng, Kích kỹ năng, Trùng sinh) mà không cần file MP3 bên ngoài.", "Bộ tổng hợp âm thanh Web Audio API:")

    add_h2("2.5. Quản lý Cơ sở dữ liệu MySQL & Server Cloud")
    add_bullet("Quản lý kết nối MySQL Pool mượt mà với fallback tự động.", "Cơ sở dữ liệu MySQL (mysql2 pool):")
    add_bullet("Triển khai trên máy chủ Cloud Railway với listener 0.0.0.0:PORT tương thích 100% môi trường Production.", "Triển khai Cloud Railway & GitHub Pages:")

    # ==========================================
    # 6. LUỒNG XỬ LÝ CHÍNH (WORKFLOWS)
    # ==========================================
    add_h1("3. LUỒNG XỬ LÝ CHÍNH (WORKFLOWS)")

    add_h2("3.1. Luồng xử lý của Người chơi Offline")
    add_p("Người chơi mở Game → Chọn Chế độ Ngoại tuyến → Chọn Mô hình Đồ họa Theme → Tiến hành Click tích điểm → Thuê công cụ nâng cấp DPC/DPS → Tích lũy sạc Nộ x2 → Kích hoạt Kỹ năng / Trùng sinh → Dữ liệu tự động lưu vào LocalStorage.")

    add_h2("3.2. Luồng xử lý của Người chơi Online Đấu trường")
    add_p("Người chơi kết nối Máy chủ → Chọn Tạo phòng Đấu trường (hoặc Nhập mã 6 số để gia nhập) → Thêm AI Bot nếu thiếu người → Bấm Sẵn sàng (Ready) → Đếm ngược 30s thi đấu Click realtime → Server tổng kết Bảng xếp hạng và trao thưởng.")

    add_h2("3.3. Luồng xử lý của Người chơi Online Hợp tác")
    add_p("Người chơi vào Phòng Hợp tác → Cùng các thành viên nhấp chuột thu gom nguyên liệu Thịt/Gỗ/Đá → Đóng góp nguyên liệu mua Nâng cấp chung → Đánh gục Boss/Cổ thụ/Vỉa đá → Hệ thống lưu vết nhật ký đồng bộ.")

    # ==========================================
    # 7. ĐẶC TẢ HỆ THỐNG CƠ SỞ DỮ LIỆU & ARCHITECTURE
    # ==========================================
    add_h1("4. ĐẶC TẢ CƠ SỞ DỮ LIỆU VÀ KIẾN TRÚC MÁY CHỦ")
    
    add_h2("4.1. Kiến trúc Cơ sở dữ liệu MySQL (Database Schema)")
    add_p("Để lưu trữ thông tin phòng chơi, tài khoản người chơi và kết quả trận đấu, hệ thống được thiết kế các bảng CSDL cốt lõi sau:")

    # Table 1: CSDL Schema Table
    table1 = doc.add_table(rows=6, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table1.rows[0].cells
    headers = ['Tên Bảng', 'Khóa Chính (PK)', 'Khóa Ngoại (FK)', 'Mô Tả Chức Năng']
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_bg(hdr_cells[idx], "7C3AED")
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    data1 = [
        ['players', 'id (VARCHAR)', 'room_code', 'Quản lý thông tin người chơi, biệt danh, chỉ số DPC, DPS, điểm số.'],
        ['rooms', 'code (VARCHAR)', '-', 'Quản lý thông tin phòng chơi online, trạng thái (lobby/playing/finished), timer.'],
        ['game_sessions', 'id (INT AUTO)', 'room_code', 'Lưu vết lịch sử các trận đấu đã hoàn thành và người chiến thắng.'],
        ['coop_resources', 'id (INT AUTO)', 'room_code', 'Lưu giữ số lượng tài nguyên Thịt, Gỗ, Đá tích lũy trong chế độ Hợp tác.'],
        ['coop_upgrades', 'id (INT AUTO)', 'room_code', 'Lưu giữ cấp độ các nâng cấp dùng chung của phòng Hợp tác.']
    ]

    for row_idx, row_data in enumerate(data1, start=1):
        row_cells = table1.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_bg(row_cells[col_idx], bg_color)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx != 1 else WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    add_p("") # spacing

    add_h2("4.2. Các tình huống phát sinh thực tế khi vận hành")
    add_bullet("Nếu thiết bị mất kết nối Socket, game tự động chuyển về chế độ Offline mà không làm gián đoạn trải nghiệm người chơi.", "Tình huống 1: Mất kết nối mạng ngột ngạt (Network Disconnect Fallback):")
    add_bullet("Chủ phòng có thể bấm Thêm Bot. Hệ thống tự động khởi tạo luồng AI Bot tự động nhấp chuột với tần số ngẫu nhiên 200-400ms.", "Tình huống 2: Thiếu người chơi trong phòng Đấu trường (AI Bot Filling):")
    add_bullet("Khi deploy lên máy chủ Railway Cloud, server tự động ràng buộc địa chỉ `0.0.0.0:${PORT}` để tiếp nhận kết nối bên ngoài mượt mà.", "Tình huống 3: Tương thích máy chủ Cloud Railway (Host Binding):")

    # ==========================================
    # 8. CÔNG NGHỆ SỬ DỤNG
    # ==========================================
    add_h1("5. CÔNG NGHỆ SỬ DỤNG IN PROJECT")

    # Tech Table
    table2 = doc.add_table(rows=5, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Phân Loại Công Nghệ"
    hdr2[1].text = "Danh Sách Công Nghệ & Thư Viện Sử Dụng"
    for cell in hdr2:
        set_cell_bg(cell, "7C3AED")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    data2 = [
        ['Frontend UI Engine', 'React 19, HTML5, Vanilla CSS3 (Glassmorphism & Sunburst animations), Lucide Icons, Web Audio API'],
        ['Backend Core Server', 'Node.js, Express.js (REST API & Static Server)'],
        ['Real-time Protocol', 'Socket.io v4 (Full-duplex WebSocket Real-time Communication)'],
        ['Database & Deployment', 'MySQL (mysql2 pool), Railway Cloud (Node Server), GitHub Pages (Static Host), GitHub Actions CI/CD']
    ]

    for r_idx, r_data in enumerate(data2, start=1):
        row_cells = table2.rows[r_idx].cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            set_cell_bg(row_cells[c_idx], bg_color)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                if c_idx == 0:
                    run.font.bold = True

    add_p("")
    add_h1("6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    add_p("Đồ án Game Siêu Clicker Tam Hợp đã hoàn thành đúng tiến độ, đáp ứng đầy đủ các yêu cầu kỹ thuật và tính năng đặt ra. Hệ thống không chỉ mang lại trải nghiệm giải trí ấn tượng với đồ họa hoạt hình 3D mượt mà mà còn khẳng định tính khả thi của việc kết hợp các công nghệ Web hiện đại (React 19, Node.js, Socket.io, MySQL) trong các bài toán thực tế.")
    add_p("Trong tương lai, nhóm sẽ tiếp tục phát triển thêm các tính năng như: Bảng xếp hạng toàn cầu (Global Leaderboard), tích hợp ví vật phẩm NFT và mở rộng thêm nhiều đồ họa Theme phong phú hơn nữa.")

    # Save to file
    out_path = 'BaoCao_GameClicker.docx'
    doc.save(out_path)
    print(f"Successfully generated {out_path}!")

if __name__ == '__main__':
    create_report()
